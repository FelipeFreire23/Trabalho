import sys
import sqlite3
from PyQt5.QtWidgets import QApplication, QWidget, QLabel, QVBoxLayout, QPushButton, QMessageBox, QTextEdit, QFileDialog
from docx import Document
from docx.shared import Inches

class CadastroDB:
    def conectar(self):
        return sqlite3.connect("cadastros.db")

    def buscar_dados(self):
        conexao = self.conectar()
        cursor = conexao.cursor()
        cursor.execute('SELECT * FROM cadastro')
        dados = cursor.fetchall()
        conexao.close()
        return dados

class Relatorio(QWidget):
    def __init__(self):
        super().__init__()
        self.db = CadastroDB()
        self.initUI()

    def initUI(self):
        self.setWindowTitle("Relatório de Cadastros")
        self.setGeometry(100, 100, 800, 600)

        layout = QVBoxLayout()

        # Área de texto para exibir o relatório
        self.text_edit = QTextEdit(self)
        self.text_edit.setReadOnly(True)
        layout.addWidget(self.text_edit)

        # Botão para gerar relatório
        self.botao_gerar = QPushButton("Gerar Relatório", self)
        self.botao_gerar.clicked.connect(self.gerar_relatorio)
        layout.addWidget(self.botao_gerar)

        # Botão para salvar como DOCX
        self.botao_docx = QPushButton("Salvar como DOCX", self)
        self.botao_docx.clicked.connect(self.salvar_docx)
        layout.addWidget(self.botao_docx)

        self.setLayout(layout)

    def gerar_relatorio(self):
        dados = self.db.buscar_dados()
        if not dados:
            QMessageBox.information(self, "Sem Dados", "Nenhum dado encontrado no banco de dados.")
            return

        # Gerar o conteúdo do relatório
        conteudo = "<h1>Relatório de Cadastros</h1>"
        conteudo += "<table border='1' cellpadding='5'><tr>" \
                    "<th>ID</th><th>Nome Completo</th><th>Data de Nascimento</th><th>Raça/Cor</th>" \
                    "<th>Escolaridade</th><th>Gênero</th><th>Telefone</th><th>Filhos</th>" \
                    "<th>Estado</th><th>Município</th></tr>"

        for linha in dados:
            conteudo += f"<tr>{''.join(f'<td>{campo}</td>' for campo in linha)}</tr>"

        conteudo += "</table>"
        self.text_edit.setHtml(conteudo)

    def salvar_docx(self):
        file_dialog = QFileDialog(self, "Salvar DOCX", "", "Documentos DOCX (*.docx)")
        if file_dialog.exec_():
            caminho_arquivo = file_dialog.selectedFiles()[0]
            if not caminho_arquivo.endswith(".docx"):
                caminho_arquivo += ".docx"

            document = Document()
            document.add_heading('Relatório de Cadastros', level=1)

            # Adiciona a tabela
            dados = self.db.buscar_dados()
            if dados:
                table = document.add_table(rows=1, cols=len(dados[0]))
                hdr_cells = table.rows[0].cells
                colunas = ["ID", "Nome Completo", "Data de Nascimento", "Raça/Cor", "Escolaridade", "Gênero", "Telefone", "Filhos", "Estado", "Município"]
                for i, coluna in enumerate(colunas):
                    hdr_cells[i].text = coluna

                for linha in dados:
                    row_cells = table.add_row().cells
                    for i, campo in enumerate(linha):
                        row_cells[i].text = str(campo)
            
            # Salva o documento
            document.save(caminho_arquivo)
            QMessageBox.information(self, "DOCX Salvo", f"Relatório salvo como DOCX em {caminho_arquivo}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    ex = Relatorio()
    ex.show()
    sys.exit(app.exec_())
